#!/usr/bin/env python3
"""Build deterministic, synthetic multi-format documents for dual-engine QA.

The Markdown files remain the human-reviewable source of truth.  The manifest
selects a representative subset of those sources as PDF, DOCX, PNG, HTML and
TXT so parser regressions are exercised without committing user data.
"""
from __future__ import annotations

import argparse
import hashlib
import html
import io
import json
import shutil
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parent
DOCUMENTS = ROOT / "documents"
FIXED_TIMESTAMP = (2026, 1, 1, 0, 0, 0)
FIXED_DATETIME = datetime(2026, 1, 1, tzinfo=timezone.utc)
IMAGE_FONT = Path("/System/Library/Fonts/STHeiti Medium.ttc")
OUTPUTS = {
    3: ("fixture_doc_03.pdf", "pdf"),
    4: ("fixture_doc_04.pdf", "pdf"),
    5: ("fixture_doc_05.docx", "docx"),
    6: ("fixture_doc_06.docx", "docx"),
    7: ("fixture_doc_07.png", "png"),
    8: ("fixture_doc_08.png", "png"),
    9: ("fixture_doc_09.html", "html"),
    10: ("fixture_doc_10.txt", "txt"),
}


def _read_source(number: int) -> tuple[str, list[str]]:
    source = DOCUMENTS / f"fixture_doc_{number:02d}.md"
    lines = source.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    paragraphs = [part.strip() for part in "\n".join(lines[1:]).split("\n\n") if part.strip()]
    if not title or not paragraphs:
        raise ValueError(f"Invalid fixture source: {source}")
    return title, paragraphs


def _wrap_by_width(text: str, *, width: float, measure) -> list[str]:
    lines: list[str] = []
    current = ""
    for character in text:
        candidate = current + character
        if current and measure(candidate) > width:
            lines.append(current)
            current = character
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def _build_pdf(destination: Path, title: str, paragraphs: list[str]) -> None:
    # Embed the fixture font so Poppler and production parsers do not depend on
    # an external Adobe-GB1 language pack.
    pdfmetrics.registerFont(TTFont("FixtureCJK", str(IMAGE_FONT), subfontIndex=0))
    page_width, page_height = letter
    document = canvas.Canvas(
        str(destination),
        pagesize=letter,
        pageCompression=1,
        invariant=1,
    )
    document.setAuthor("GraphRAG Studio synthetic evaluation")
    document.setCreator("GraphRAG Studio fixture generator")
    document.setTitle(title)
    left = right = 72
    y = page_height - 76
    document.setFillColorRGB(46 / 255, 116 / 255, 181 / 255)
    document.setFont("FixtureCJK", 20)
    document.drawString(left, y, title)
    y -= 38
    document.setFillColorRGB(0.12, 0.14, 0.17)
    document.setFont("FixtureCJK", 11)
    for paragraph in paragraphs:
        for line in _wrap_by_width(
            paragraph,
            width=page_width - left - right,
            measure=lambda value: pdfmetrics.stringWidth(value, "FixtureCJK", 11),
        ):
            document.drawString(left, y, line)
            y -= 18
        y -= 10
    document.setFillColorRGB(0.45, 0.47, 0.50)
    document.setFont("FixtureCJK", 9)
    document.drawRightString(page_width - right, 40, "合成评测语料 · 第 1 页")
    document.showPage()
    document.save()


def _set_run_font(run, *, name: str, east_asia: str, size: int, color: str | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _normalize_docx_archive(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source:
        entries = [(info.filename, source.read(info.filename)) for info in source.infolist()]
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for filename, content in sorted(entries):
            info = zipfile.ZipInfo(filename, FIXED_TIMESTAMP)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, content)
    path.write_bytes(buffer.getvalue())


def _build_docx(destination: Path, title: str, paragraphs: list[str]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # Named CJK override for deterministic macOS/LibreOffice fixture rendering.
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(46, 116, 181)
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)

    title_paragraph = document.add_paragraph(style="Heading 1")
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run, name="Calibri", east_asia="Heiti SC", size=16, color="2E74B5")
    for text in paragraphs:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        _set_run_font(run, name="Calibri", east_asia="Heiti SC", size=11)

    properties = document.core_properties
    properties.author = "GraphRAG Studio synthetic evaluation"
    properties.creator = "GraphRAG Studio fixture generator"
    properties.created = FIXED_DATETIME
    properties.modified = FIXED_DATETIME
    properties.title = title
    document.save(destination)
    _normalize_docx_archive(destination)


def _font(size: int) -> ImageFont.FreeTypeFont:
    if not IMAGE_FONT.is_file():
        raise FileNotFoundError(f"Chinese fixture font is unavailable: {IMAGE_FONT}")
    return ImageFont.truetype(str(IMAGE_FONT), size=size)


def _build_png(destination: Path, title: str, paragraphs: list[str]) -> None:
    image = Image.new("RGB", (1600, 1000), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    title_font = _font(48)
    body_font = _font(28)
    footer_font = _font(20)
    draw.rounded_rectangle((80, 72, 1520, 928), radius=24, fill="#FFFFFF", outline="#D8E0EA", width=3)
    draw.text((132, 122), title, font=title_font, fill="#2E74B5")
    y = 214
    for paragraph in paragraphs:
        lines = _wrap_by_width(
            paragraph,
            width=1328,
            measure=lambda value: draw.textlength(value, font=body_font),
        )
        for line in lines:
            draw.text((132, y), line, font=body_font, fill="#20242B")
            y += 48
        y += 24
    draw.text((132, 866), "脱敏合成图片语料 · 仅用于解析回归", font=footer_font, fill="#6B7280")
    image.save(destination, format="PNG", optimize=True)


def _build_html(destination: Path, title: str, paragraphs: list[str]) -> None:
    body = "\n".join(f"    <p>{html.escape(paragraph)}</p>" for paragraph in paragraphs)
    destination.write_text(
        "<!doctype html>\n"
        '<html lang="zh-CN">\n'
        "  <head>\n"
        '    <meta charset="utf-8">\n'
        f"    <title>{html.escape(title)}</title>\n"
        "  </head>\n"
        "  <body>\n"
        f"    <h1>{html.escape(title)}</h1>\n"
        f"{body}\n"
        "  </body>\n"
        "</html>\n",
        encoding="utf-8",
    )


def _build_txt(destination: Path, title: str, paragraphs: list[str]) -> None:
    destination.write_text(f"{title}\n\n" + "\n\n".join(paragraphs) + "\n", encoding="utf-8")


BUILDERS = {
    "pdf": _build_pdf,
    "docx": _build_docx,
    "png": _build_png,
    "html": _build_html,
    "txt": _build_txt,
}


def build(output_dir: Path) -> dict[str, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    hashes: dict[str, str] = {}
    for number, (filename, kind) in OUTPUTS.items():
        title, paragraphs = _read_source(number)
        destination = output_dir / filename
        BUILDERS[kind](destination, title, paragraphs)
        hashes[filename] = hashlib.sha256(destination.read_bytes()).hexdigest()
    return hashes


def check() -> None:
    with tempfile.TemporaryDirectory(prefix="graphrag-eval-fixtures-") as directory:
        expected = build(Path(directory))
        actual = {
            filename: hashlib.sha256((DOCUMENTS / filename).read_bytes()).hexdigest()
            for filename, _kind in OUTPUTS.values()
        }
    mismatches = [filename for filename, digest in expected.items() if actual.get(filename) != digest]
    if mismatches:
        raise SystemExit("Fixture outputs are stale: " + ", ".join(mismatches))
    print(json.dumps({"ok": True, "files": len(actual), "sha256": actual}, ensure_ascii=False, indent=2))


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--check", action="store_true", help="verify committed outputs are reproducible")
    args = parser.parse_args()
    if args.check:
        check()
        return
    hashes = build(DOCUMENTS)
    print(json.dumps({"generated": len(hashes), "sha256": hashes}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
